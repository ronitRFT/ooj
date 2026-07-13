#!/usr/bin/env python3
"""Generate OOJ Foundation documentation as Word .docx files."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import re

DOCS_DIR = Path(__file__).parent
OUTPUT_DIR = DOCS_DIR


def add_title_page(doc, title, subtitle, meta_lines):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x8B, 0x45, 0x13)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(subtitle)
    run2.font.size = Pt(14)
    run2.italic = True

    doc.add_paragraph()
    for line in meta_lines:
        mp = doc.add_paragraph()
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mr = mp.add_run(line)
        mr.font.size = Pt(10)

    doc.add_page_break()


def parse_markdown_to_docx(md_path: Path, docx_path: Path, title: str, subtitle: str, meta: list):
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    add_title_page(doc, title, subtitle, meta)

    content = md_path.read_text(encoding='utf-8')
    lines = content.split('\n')

    i = 0
    in_code = False
    code_buffer = []

    while i < len(lines):
        line = lines[i]

        # Skip first title block already on cover
        if i < 5 and line.startswith('# '):
            i += 1
            continue

        if line.strip().startswith('```'):
            if in_code:
                p = doc.add_paragraph('\n'.join(code_buffer))
                p.style = 'No Spacing'
                for run in p.runs:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                code_buffer = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buffer.append(line)
            i += 1
            continue

        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=4)
        elif line.startswith('---'):
            doc.add_paragraph('─' * 60)
        elif line.startswith('|') and '|' in line[1:]:
            # table row
            cells = [c.strip() for c in line.strip('|').split('|')]
            if all(set(c) <= set('-: ') for c in cells):
                i += 1
                continue  # skip separator
            # collect table rows
            table_rows = [cells]
            j = i + 1
            while j < len(lines) and lines[j].startswith('|'):
                row_cells = [c.strip() for c in lines[j].strip('|').split('|')]
                if not all(set(c) <= set('-: ') for c in row_cells):
                    table_rows.append(row_cells)
                j += 1
            if table_rows:
                cols = max(len(r) for r in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=cols)
                table.style = 'Table Grid'
                for ri, row in enumerate(table_rows):
                    for ci, cell in enumerate(row):
                        if ci < cols:
                            table.rows[ri].cells[ci].text = cell
                doc.add_paragraph()
            i = j
            continue
        elif line.startswith('- '):
            doc.add_paragraph(line[2:].strip(), style='List Bullet')
        elif re.match(r'^\d+\.\s', line):
            doc.add_paragraph(re.sub(r'^\d+\.\s', '', line).strip(), style='List Number')
        elif line.strip() == '':
            pass
        else:
            # bold inline
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

        i += 1

    doc.save(docx_path)
    print(f'Created: {docx_path}')


if __name__ == '__main__':
    parse_markdown_to_docx(
        DOCS_DIR / 'OOJ-Technical-Vision-Document-v1.md',
        OUTPUT_DIR / 'OOJ-Technical-Vision-Document-v1.docx',
        'OOJ FOUNDATION EVENT MANAGEMENT SYSTEM',
        'Technical Vision & Product Architecture Document — v1.0',
        [
            'Document Status: As-Built — Reflects Current Codebase',
            'API Version: v1 (REST) — base path /api',
            'Tech Stack: Node.js · Express · React 19 · Vite · MySQL 8',
            'Prepared For: OOJ Foundation / Engineering & Operations Review',
            'Classification: Internal Use — Foundation Operations',
        ],
    )

    parse_markdown_to_docx(
        DOCS_DIR / 'OOJ-Swimlane-Workflow-Document.md',
        OUTPUT_DIR / 'OOJ-Swimlane-Workflow-Document.docx',
        'OOJ FOUNDATION EVENT MANAGEMENT SYSTEM',
        'Swimlane Workflow Diagrams & RACI Matrices',
        [
            'RACI: R = Responsible · A = Accountable · C = Consulted · I = Informed',
            'Document Status: As-Built — Reflects Current Codebase',
            'Prepared For: OOJ Foundation / Operations Review',
            'Classification: Internal Use — Foundation Operations',
        ],
    )
